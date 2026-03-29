#!/usr/bin/env python3
"""Build Word (.docx) CEO packages for Eyal: executive summary, site map, decisions.

Outputs go to docs/project/eyal-ceo-submissions-and-responses/to-eyal/
and executive is also mirrored to team-100-preplanning/ for scripts compatibility.

Requires: pip install python-docx
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx_rtl import add_md_file_as_docx, add_rtl_paragraph, set_run_font


DELIVERY_DATE = "2026-03-29"
VERSION_TAG = "v1"


def _paths():
    root = Path(__file__).resolve().parents[1]
    to_eyal = root / "docs/project/eyal-ceo-submissions-and-responses/to-eyal"
    team100 = root / "docs/project/team-100-preplanning"
    return root, to_eyal, team100


def build_executive_summary_docx():
    _, to_eyal, team100 = _paths()
    out_primary = to_eyal / f"{DELIVERY_DATE}--executive-summary--{VERSION_TAG}.docx"
    out_mirror = team100 / "EYAL-EXECUTIVE-SUMMARY-FOR-EYAL.docx"

    doc = Document()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = t.add_run("תקציר מנהלים + טופס אישור — אייל עמית")
    set_run_font(r, size=18, bold=True)

    add_rtl_paragraph(doc, f"גרסה: 1.0  |  תאריך: {DELIVERY_DATE}")
    add_rtl_paragraph(doc, "הוגש על ידי: נימרוד (ליווי מיגרציה)")
    add_rtl_paragraph(doc, "סטטוס: ממתין חתימת אייל עמית", bold=True)
    doc.add_paragraph()

    add_rtl_paragraph(doc, "תקציר מנהלים", heading_level=1)

    body = """מטרה: לצאת מלופי התיקון של האתר הישן ולבנות אתר WordPress אחד, פשוט לתחזוקה, עם דגש על מבקרים, SEO, ונגישות מלאה לפי דרישות החוק בישראל. המכירות והסליקה יבוצעו מחוץ לאתר באמצעות חשבונית ירוקה; באתר יוצגו עמודי מוצר/שירות פשוטים עם כפתור להפניה לסליקה.

החלטות שאושרו בתכנון (עם נימרוד) — לאישורך הסופי:

1. פלטפורמה: WordPress; עבודה בענף פיתוח נקי; אפשרות להסתמך על האתר החי כמקור אמת לתוכן ולהעביר את סבב הפיתוח הקודם לארכיון.

2. מותגים: מגוון מותגים באותו דומיין — הפרדה בחוויית משתמש בלבד (לא מערכות נפרדות).

3. מרכז: הסטודיו, הנשימה והדיג'רידו — המרכז בדף הבית ובתפריט הראשי.

4. הוצאה ומופע: חלק מהעבר, נשארים וחשובים — מוצגים כארכיון מותגי משני שמעשיר את האתר.

5. חנות: ללא עגלה ותשלום באתר; עמודים פשוטים + חשבונית ירוקה לסליקה.

6. עמודי QR: חובה לשמור את כל הכתובות (מודפסים בספרים). אסור לשנות slug או לבצע 301 ששובר סריקה. יתווסף עמוד אינדקס (ציבורי או למנהלים — לבחירתך).

7. בלוג: להחזיר לחיים כנכס SEO מרכזי.

8. אירועים: בלוק "אירוע הבא" בתבנית; טפסים — קצר באתר + מפורט חיצונית.

9. SEO מקומי: סכמות עסק מקומי (לפי מה שמאושר מקצועית).

10. נגישות: עמידה מלאה בדרישות החוק — כולל מסמך LEGAL-ACCESSIBILITY-ISRAEL-SPEC במאגר (המלצה לייעוץ חיצוני).

11. תפעול: אייל — מנהל ומתחזק; נימרוד — ליווי עד יציבות.

12. שפה: עמוד נחיתה באנגלית בנוסף לעברית.

מדדי הצלחה: לאזן יציבות, תחזוקה, SEO ו-UX — וגם תנועה אמיתית; בלי מבקרים היעד העסקי לא הושג."""

    for para in body.split("\n\n"):
        add_rtl_paragraph(doc, para.strip())

    doc.add_paragraph()
    add_rtl_paragraph(doc, "אישור עקרונות (סמן וחתום)", heading_level=1)

    table = doc.add_table(rows=8, cols=3)
    table.style = "Table Grid"
    hdr = ("#", "נושא", "מאושר (סמן V)")
    for i, h in enumerate(hdr):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(h)
        set_run_font(run, bold=True)
    rows_data = [
        ("1", "אתר WordPress אחד, ענף נקי, ארכיון מסלול קודם", ""),
        ("2", "הפרדת UX בין סטודיו / הוצאה+ספרים / מופעים", ""),
        ("3", "סליקה בחשבונית ירוקה בלבד (ללא עגלה באתר)", ""),
        ("4", "שימור מלא URLי QR ללא שבירת ספרים מודפסים", ""),
        ("5", "בלוג כנכס SEO מחודש", ""),
        ("6", "נגישות — עמידה מלאה + בחינת ייעוץ חיצוני לפי תקציב", ""),
        ("7", "עמוד נחיתה באנגלית", ""),
    ]
    for r_idx, (num, topic, _) in enumerate(rows_data, start=1):
        row = table.rows[r_idx].cells
        for c_idx, val in enumerate((num, topic, "_______")):
            p = row[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(val)
            set_run_font(run)

    doc.add_paragraph()
    add_rtl_paragraph(doc, "שם: _________________________    חתימה: __________________    תאריך: __________")
    doc.add_paragraph()
    add_rtl_paragraph(doc, "הערות אייל:", bold=True)
    add_rtl_paragraph(doc, "_________________________________________________________________")
    add_rtl_paragraph(doc, "_________________________________________________________________")

    doc.add_paragraph()
    add_rtl_paragraph(doc, "פתוח למילוי על ידי אייל (לפני בנייה מלאה)", heading_level=1)

    open_items = """1. חשבונית ירוקה: רשימת מוצרים/שירותים וקישורי סליקה לדוגמה לכל סוג — ראו במאגר: GREEN-INVOICE-LINK-MAP.md

2. אינדקס QR: ציבורי או רק למנהל מחובר? רמת פירוט (רשימה / חיפוש).

3. שמות תפריט בעברית לשלושת המרחבים (סטודיו / הוצאה / מופע).

4. עמוד EN: מטרה (תיירים / בינלאומי) וטקסט מקור או הנחיה לכותב.

5. ייעוץ נגישות חיצוני: מאושר תקציב כן/לא.

6. 301 לעמודי /shop/ — יעד מדויק (צור קשר / דף מידע) אחרי בדיקה שאין קישורים קריטיים.

לאחר מילוי: להחזיר לנימרוד בערוץ העבודה."""

    for para in open_items.split("\n\n"):
        add_rtl_paragraph(doc, para.strip())

    doc.add_paragraph()
    add_rtl_paragraph(
        doc,
        "קבצים במאגר: eyalamit.co.il/docs/project/eyal-ceo-submissions-and-responses/to-eyal/ (הגשה) + team-100-preplanning/ (מקורות Markdown).",
        size=9,
    )

    to_eyal.mkdir(parents=True, exist_ok=True)
    doc.save(out_primary)
    doc.save(out_mirror)
    return out_primary, out_mirror


def build_site_map_docx():
    _, to_eyal, team100 = _paths()
    out = to_eyal / f"{DELIVERY_DATE}--site-map-draft-v2--{VERSION_TAG}.docx"
    md_path = team100 / "SITEMAP-NEW-SITE-v2-DRAFT.md"

    doc = Document()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = t.add_run("מפת אתר — טיוטה לאישור אייל (v2)")
    set_run_font(r, size=18, bold=True)
    add_rtl_paragraph(doc, f"תאריך הגשה: {DELIVERY_DATE}  |  סטטוס: טיוטה — ממתין אישור")
    add_rtl_paragraph(
        doc,
        "מסמך זה משקף את מבנה המידע המוצע לאתר החדש. אחרי אישורך נעבור לאפיון עמודים מפורט.",
        size=10,
    )
    doc.add_paragraph()
    add_md_file_as_docx(doc, md_path, skip_first_h1=True)
    doc.add_paragraph()
    add_rtl_paragraph(doc, "חתימה לאישור מפת אתר: __________________  תאריך: __________", bold=True)
    to_eyal.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def _add_table(doc, headers: tuple, rows: tuple):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        p = table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(h)
        set_run_font(run, bold=True)
    for ri, row_vals in enumerate(rows, start=1):
        for ci, val in enumerate(row_vals):
            p = table.rows[ri].cells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(val)
            set_run_font(run)
    return table


def build_decisions_docx():
    _, to_eyal, team100 = _paths()
    out = to_eyal / f"{DELIVERY_DATE}--decisions-for-approval--{VERSION_TAG}.docx"

    doc = Document()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = t.add_run("קובץ החלטות לאישור אייל — אתר חדש eyalamit.co.il")
    set_run_font(r, size=18, bold=True)
    add_rtl_paragraph(doc, f"גרסה: 2.0 (סיכום)  |  תאריך: {DELIVERY_DATE}")
    add_rtl_paragraph(doc, "מקור מפורט בצוות: שער 0, מדיניות QR, טבלת Keep/Merge/Drop מול מאגר התוכן.", size=10)
    doc.add_paragraph()

    add_rtl_paragraph(doc, "1. מהות האתר (שער 0)", heading_level=1)
    _add_table(
        doc,
        ("הצהרה", "נדרש אישור (סמן)"),
        (
            ("האתר הוא אתר שיווקי־מותגי: מידע, אמון, והנעה לפעולה", "[ ]"),
            ("האתר אינו מערכת מכירות מקוונת או עגלת קניות", "[ ]"),
            ("מורכבות (מכירה, תשלומים) — במערכות חיצוניות לפי צורך", "[ ]"),
        ),
    )
    doc.add_paragraph()
    add_rtl_paragraph(doc, "חתימה: __________________  תאריך: ________")

    doc.add_paragraph()
    add_rtl_paragraph(doc, "2. מדדי הצלחה (תמצית)", heading_level=1)
    _add_table(
        doc,
        ("מדד", "יעד מוצע"),
        (
            ("ביצועים מובייל", "Lighthouse Performance ≥ 85 (או שקיל)"),
            ("יציבות", "0 שגיאות קונסול קריטיות בעמודי ליבה"),
            ("תחזוקה", "רשימת תוספים מאושרת ומצומצמת"),
            ("SEO טכני", "sitemap אחד, ללא כפילויות קנוניקל חמורות"),
            ("בהירות מסר", "מבקר מבין מה נעשה, למי, איך מתחילים"),
        ),
    )
    add_rtl_paragraph(doc, "תקרת תקציב שנתית (למילוי): __________________")

    doc.add_paragraph()
    add_rtl_paragraph(doc, "3. כיוון טכני (סמן אפשרות אחת)", heading_level=1)
    _add_table(
        doc,
        ("אפשרות", "תיאור", "סמן"),
        (
            ("A", "מופע WordPress חדש ורזה; המונולית כארכיון הפניה", "[ ]"),
            ("B", "אתר סטטי / Jamstack + CMS קל", "[ ]"),
            ("C", "המשך ייצוב המונולית — דורש נימוק בכתב", "[ ]"),
        ),
    )
    add_rtl_paragraph(doc, "המלצת צוות 100: אפשרות A.")
    add_rtl_paragraph(doc, "החלטה סופית (A / B / C): ________  חתימה: __________________")

    doc.add_paragraph()
    add_rtl_paragraph(doc, "4. החלטות מתועדות (לסיכום — לאישורך)", heading_level=1)
    bullets = (
        "מותגים: אתר אחד, דומיין אחד; הפרדת UX בין סטודיו (מרכז), הוצאה/ספרים, מופעים (ארכיון).",
        "מקור אמת לתוכן: אתר חי + ענף Git; מסלול מונולית קודם — ארכיון.",
        "מסחר: ללא עגלה ב-WordPress; סליקה בחשבונית ירוקה; עמודי מוצר פשוטים.",
        "QR: שימור URL קבוע — לפי מדיניות QR במאגר (ספרים מודפסים).",
        "בלוג: החייאה כנכס SEO.",
        "נגישות: עמידה מלאה + בחינת ייעוץ חיצוני לפי תקציב.",
        "תפעול: אייל — מנהל; נימרוד — מיגרציה עד יציבות.",
        "אנגלית: עמוד נחיתה באנגלית.",
    )
    for b in bullets:
        add_rtl_paragraph(doc, "• " + b)

    doc.add_paragraph()
    add_rtl_paragraph(doc, "5. מדיניות תוכן (Keep / Merge / Drop) — עקרונות מחייבים", heading_level=1)
    add_rtl_paragraph(
        doc,
        "• QR: Keep לנתיב — אין שינוי slug ואין 301 (אלא אם תחליט אחרת בכתב).",
    )
    add_rtl_paragraph(
        doc,
        "• חנות Woo: Drop ממבנה האתר החדש + 301 ליעד שתאשר (לאחר בדיקת קישורים).",
    )
    add_rtl_paragraph(
        doc,
        "• בלוג: לפי כוונתך (החייאה ואופטימיזציה); סינון תוכן חלש — המלצה משנית.",
    )
    add_rtl_paragraph(
        doc,
        "רשימת כל העמודים והפוסטים (שורה-שורה) נמצאת בקובץ CONTENT-SSOT-INVENTORY במאגר — תמולא במסגרת הפיתוח לאחר אישור עקרונות זה.",
        size=10,
    )

    doc.add_paragraph()
    add_rtl_paragraph(doc, "הערות אייל:", bold=True)
    add_rtl_paragraph(doc, "_________________________________________________________________")
    doc.add_paragraph()
    add_rtl_paragraph(doc, "מסמך מקור Markdown לצוות: team-100-preplanning/01-GATE-ZERO-STRATEGY.md וכו'.", size=9)

    to_eyal.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def main():
    _, to_eyal, _ = _paths()
    e1, e2 = build_executive_summary_docx()
    s = build_site_map_docx()
    d = build_decisions_docx()
    print("Written:")
    for p in (e1, e2, s, d):
        print(" ", p)


if __name__ == "__main__":
    main()
